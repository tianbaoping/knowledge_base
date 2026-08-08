import os
import re
import json
import uuid
import hashlib
import time
from typing import List, Dict, Any, Optional, Tuple, Callable
from datetime import datetime
from loguru import logger
from app.config import settings
from app.services.ocr_service import ocr_service_monitor


class ParseProgress:
    """解析进度追踪器"""

    STAGE_INIT = "init"
    STAGE_DETECT = "detect"
    STAGE_EXTRACT = "extract"
    STAGE_OCR = "ocr"
    STAGE_CHUNK = "chunk"
    STAGE_COMPLETE = "complete"
    STAGE_ERROR = "error"

    def __init__(self, file_name: str):
        self.file_name = file_name
        self.stage = self.STAGE_INIT
        self.progress_pct = 0
        self.message = "初始化中..."
        self.needs_ocr = False
        self.ocr_reason = ""
        self.stages_completed = []
        self.start_time = time.time()
        self.elapsed_time = 0
        self.error = None

    def update(self, stage: str, progress_pct: float, message: str,
               needs_ocr: bool = False, ocr_reason: str = ""):
        self.stage = stage
        self.progress_pct = min(progress_pct, 100)
        self.message = message
        if needs_ocr:
            self.needs_ocr = True
            self.ocr_reason = ocr_reason
        if stage not in self.stages_completed:
            self.stages_completed.append(stage)
        self.elapsed_time = time.time() - self.start_time

    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_name": self.file_name,
            "stage": self.stage,
            "progress": self.progress_pct,
            "message": self.message,
            "needs_ocr": self.needs_ocr,
            "ocr_reason": self.ocr_reason,
            "stages_completed": self.stages_completed,
            "elapsed_time": round(self.elapsed_time, 2),
            "error": self.error,
        }


class DocumentParser:
    """
    文档解析器

    解析管道流程:
    1. detect: 检测文件类型，判断是否需要 OCR
    2. extract: 提取文本（直接提取或 OCR）
    3. chunk: 切片处理
    4. complete: 完成

    支持的文件类型:
    - 文本类: txt, md
    - Word: doc, docx
    - PDF: 自动检测是否为扫描版，扫描版使用 OCR
    - OFD: 中国版式文档标准 (GB/T 33190)
    - 图片: jpg, jpeg, png, bmp, tiff, gif (使用 OCR)
    """

    # 直接文本提取的格式
    TEXT_FORMATS = {"txt", "md"}
    WORD_FORMATS = {"doc", "docx"}
    PDF_FORMATS = {"pdf"}
    OFD_FORMATS = {"ofd"}
    IMAGE_FORMATS = {"jpg", "jpeg", "png", "bmp", "tiff", "gif"}

    def __init__(self):
        self.supported_formats = settings.SUPPORTED_FORMATS
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self._ocr_engine = None

    def set_ocr_engine(self, ocr_engine):
        """设置 OCR 引擎引用"""
        self._ocr_engine = ocr_engine

    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        if not os.path.exists(file_path):
            return False, "文件不存在"
        if not os.path.isfile(file_path):
            return False, "路径不是文件"
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "文件为空"
        if file_size > settings.MAX_FILE_SIZE:
            return False, f"文件超过{settings.MAX_FILE_SIZE // 1024 // 1024}MB限制"
        ext = self._get_extension(file_path)
        if ext not in self.supported_formats:
            return False, f"不支持的文件格式: {ext}"
        if self._is_encrypted(file_path):
            return False, "文件已加密，无法解析"
        return True, "文件校验通过"

    def _get_extension(self, file_path: str) -> str:
        return os.path.splitext(file_path)[1].lower().lstrip(".")

    def _is_encrypted(self, file_path: str) -> bool:
        ext = self._get_extension(file_path)
        try:
            if ext == "pdf":
                import fitz
                doc = fitz.open(file_path)
                is_encrypted = doc.is_encrypted
                doc.close()
                return is_encrypted
            elif ext == "docx":
                # .docx 是 ZIP 格式，检查是否可读
                import zipfile
                if zipfile.is_zipfile(file_path):
                    return False
                return True
            elif ext == "doc":
                # .doc 是 OLE 二进制格式，尝试用 olefile 检测加密
                try:
                    import olefile
                    if olefile.isOleFile(file_path):
                        ole = olefile.OleFileIO(file_path)
                        ole.close()
                        return False
                except Exception:
                    pass
                # 如果 olefile 不可用，尝试读取文件头判断
                with open(file_path, "rb") as f:
                    magic = f.read(8)
                    # OLE 文件魔数: D0 CF 11 E0 A1 B1 1A E1
                    if magic[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                        return False
                return True
            return False
        except Exception:
            return True

    def compute_md5(self, file_path: str) -> str:
        md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                md5.update(chunk)
        return md5.hexdigest()

    def detect_file_type(self, file_path: str) -> Dict[str, Any]:
        """
        检测文件类型和是否需要 OCR

        Returns:
            {
                "file_type": "text" | "word" | "pdf_text" | "pdf_scanned" | "ofd" | "image",
                "ext": "pdf",
                "needs_ocr": bool,
                "reason": str,
            }
        """
        ext = self._get_extension(file_path)

        if ext in self.TEXT_FORMATS:
            return {
                "file_type": "text",
                "ext": ext,
                "needs_ocr": False,
                "reason": "文本文件，直接解析",
            }

        elif ext in self.WORD_FORMATS:
            return {
                "file_type": "word",
                "ext": ext,
                "needs_ocr": False,
                "reason": "Word 文件，直接解析",
            }

        elif ext in self.PDF_FORMATS:
            # 检测 PDF 是否为扫描版
            return self._detect_pdf_type(file_path)

        elif ext in self.OFD_FORMATS:
            return {
                "file_type": "ofd",
                "ext": ext,
                "needs_ocr": False,
                "reason": "OFD 版式文档，使用 OFDParser 直接解析",
            }

        elif ext in self.IMAGE_FORMATS:
            return {
                "file_type": "image",
                "ext": ext,
                "needs_ocr": True,
                "reason": "图片文件，需要 OCR 识别",
            }

        else:
            return {
                "file_type": "unknown",
                "ext": ext,
                "needs_ocr": False,
                "reason": f"未知格式: {ext}",
            }

    def _detect_pdf_type(self, file_path: str) -> Dict[str, Any]:
        """
        检测 PDF 类型：文本 PDF vs 扫描版 PDF

        策略:
        - 检查前 3 页的文本含量
        - 如果文本极少（平均每页 < 50 字符），判定为扫描版
        - 如果有可提取文本，判定为文本 PDF
        """
        try:
            import fitz
            doc = fitz.open(file_path)

            if doc.is_encrypted:
                doc.close()
                return {
                    "file_type": "pdf_scanned",
                    "ext": "pdf",
                    "needs_ocr": True,
                    "reason": "PDF 已加密，需要 OCR 处理",
                }

            total_pages = len(doc)
            check_pages = min(3, total_pages)
            total_text_len = 0
            has_images = False

            for i in range(check_pages):
                page = doc[i]
                text = page.get_text().strip()
                total_text_len += len(text)

                # 检查是否有图片
                images = page.get_images(full=True)
                if len(images) > 0:
                    has_images = True

            doc.close()

            avg_text_len = total_text_len / check_pages if check_pages > 0 else 0

            # 判断标准
            if avg_text_len < 50:
                # 文本极少，可能是扫描版
                return {
                    "file_type": "pdf_scanned",
                    "ext": "pdf",
                    "needs_ocr": True,
                    "reason": f"PDF 文本含量极少（平均每页 {avg_text_len:.0f} 字符），判定为扫描版，需要 OCR",
                }
            else:
                return {
                    "file_type": "pdf_text",
                    "ext": "pdf",
                    "needs_ocr": False,
                    "reason": f"PDF 含可提取文本（平均每页 {avg_text_len:.0f} 字符），直接解析",
                }

        except ImportError:
            return {
                "file_type": "pdf_scanned",
                "ext": "pdf",
                "needs_ocr": True,
                "reason": "缺少 PyMuPDF，默认使用 OCR",
            }
        except Exception as e:
            logger.warning(f"PDF 类型检测失败: {e}，默认使用 OCR")
            return {
                "file_type": "pdf_scanned",
                "ext": "pdf",
                "needs_ocr": True,
                "reason": f"PDF 类型检测失败: {e}",
            }

    def parse(self, file_path: str, progress_callback: Optional[Callable] = None) -> Tuple[Optional[Dict[str, Any]], str]:
        """
        解析文件

        Args:
            file_path: 文件路径
            progress_callback: 进度回调函数 (progress: ParseProgress) -> None

        Returns:
            (parsed_data, message)
        """
        file_name = os.path.basename(file_path)
        progress = ParseProgress(file_name)

        try:
            # Stage 1: 检测文件类型
            progress.update(ParseProgress.STAGE_DETECT, 10, "检测文件类型...")
            self._notify_progress(progress, progress_callback)

            file_info = self.detect_file_type(file_path)
            progress.needs_ocr = file_info["needs_ocr"]
            progress.ocr_reason = file_info["reason"]

            if file_info["needs_ocr"]:
                progress.update(
                    ParseProgress.STAGE_DETECT, 20,
                    f"检测为需要 OCR 处理: {file_info['reason']}"
                )
            else:
                progress.update(
                    ParseProgress.STAGE_DETECT, 20,
                    f"检测为直接解析: {file_info['reason']}"
                )
            self._notify_progress(progress, progress_callback)

            # Stage 2: 提取文本
            progress.update(ParseProgress.STAGE_EXTRACT, 30, "提取文本内容...")
            self._notify_progress(progress, progress_callback)

            file_type = file_info["file_type"]
            result_data = None
            result_msg = ""

            if file_type == "text":
                result_data, result_msg = self._parse_txt(file_path)
            elif file_type == "word":
                result_data, result_msg = self._parse_word(file_path)
            elif file_type == "ofd":
                result_data, result_msg = self._parse_ofd(file_path)
            elif file_type == "pdf_text":
                result_data, result_msg = self._parse_pdf(file_path)
            elif file_type == "pdf_scanned":
                progress.update(ParseProgress.STAGE_OCR, 40, "使用 OCR 识别扫描版 PDF...")
                self._notify_progress(progress, progress_callback)
                result_data = self._parse_pdf_with_ocr(file_path, progress, progress_callback)
                result_msg = "PDF OCR 解析成功" if result_data else "PDF OCR 解析失败"
            elif file_type == "image":
                progress.update(ParseProgress.STAGE_OCR, 40, "使用 OCR 识别图片...")
                self._notify_progress(progress, progress_callback)
                result_data = self._parse_image_with_ocr(file_path)
                result_msg = "图片 OCR 解析成功" if result_data else "图片 OCR 解析失败"
            else:
                return None, f"暂不支持的文件类型: {file_type}"

            if result_data is None:
                progress.error = result_msg or "解析失败"
                progress.update(ParseProgress.STAGE_ERROR, 100, f"解析失败: {progress.error}")
                self._notify_progress(progress, progress_callback)
                return None, progress.error

            progress.update(ParseProgress.STAGE_EXTRACT, 70, "文本提取完成")
            self._notify_progress(progress, progress_callback)

            # Stage 3: 完成
            progress.update(ParseProgress.STAGE_COMPLETE, 100, "解析完成")
            self._notify_progress(progress, progress_callback)

            return result_data, "解析成功"

        except Exception as e:
            logger.error(f"文件解析异常 {file_path}: {e}")
            progress.error = str(e)
            progress.update(ParseProgress.STAGE_ERROR, 100, f"解析异常: {str(e)}")
            self._notify_progress(progress, progress_callback)
            return None, f"解析异常: {str(e)}"

    def _notify_progress(self, progress: ParseProgress, callback: Optional[Callable]):
        """通知进度回调"""
        if callback:
            try:
                callback(progress)
            except Exception as e:
                logger.debug(f"进度回调失败: {e}")

    def _parse_pdf_with_ocr(self, file_path: str, progress: ParseProgress = None,
                            progress_callback: Optional[Callable] = None) -> Optional[Dict[str, Any]]:
        """
        使用 OCR 解析 PDF（扫描版）
        """
        if self._ocr_engine is None:
            logger.warning("OCR 引擎未初始化，尝试直接解析")
            result, msg = self._parse_pdf(file_path)
            if result:
                result["parse_method"] = "pdf_text_fallback"
                result["ocr_fallback"] = True
            return result

        start_time = time.time()
        try:
            import fitz

            with open(file_path, "rb") as f:
                pdf_bytes = f.read()

            # 获取 PDF 基本信息
            doc = fitz.open(file_path)
            total_pages = len(doc)
            doc.close()

            # 使用 OCR 引擎识别
            ocr_result = self._ocr_engine.recognize_pdf(
                pdf_bytes,
                max_pages=50,
                dpi=200,
            )

            # 转换为统一格式
            pages_text = []
            for page_data in ocr_result.get("pages", []):
                page_num = page_data["page"]
                text = page_data.get("text", "")
                if text.strip():
                    pages_text.append({
                        "page": page_num,
                        "text": text.strip(),
                        "ocr_lines": page_data.get("lines", []),
                    })

            full_text = ocr_result.get("full_text", "")
            full_text = self._clean_text(full_text)

            processing_time_ms = int((time.time() - start_time) * 1000)

            # 记录 OCR 请求统计
            ocr_service_monitor.record_request(
                pages=total_pages,
                chars=len(full_text),
                processing_time_ms=processing_time_ms,
                success=True,
            )

            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "pdf",
                "parse_method": "ocr",
                "is_scanned": True,
                "total_pages": total_pages,
                "text_content": full_text,
                "pages": pages_text,
                "charts": [],
                "chart_count": 0,
                "char_count": len(full_text),
                "ocr_result": ocr_result,
            }

            logger.info(f"PDF OCR 解析成功: {file_path}, 共{total_pages}页, 耗时 {processing_time_ms}ms")
            return result

        except ImportError:
            logger.error("缺少 PyMuPDF 依赖")
            return None
        except Exception as e:
            processing_time_ms = int((time.time() - start_time) * 1000)
            ocr_service_monitor.record_request(
                pages=0,
                chars=0,
                processing_time_ms=processing_time_ms,
                success=False,
            )
            logger.error(f"PDF OCR 解析失败: {e}")
            return None

    def _parse_image_with_ocr(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        使用 OCR 解析图片
        """
        if self._ocr_engine is None:
            logger.warning("OCR 引擎未初始化，无法解析图片")
            return None

        start_time = time.time()
        try:
            with open(file_path, "rb") as f:
                image_bytes = f.read()

            ocr_result = self._ocr_engine.recognize_image(image_bytes, page_num=1)

            full_text = ocr_result.get("text", "")
            full_text = self._clean_text(full_text)

            processing_time_ms = int((time.time() - start_time) * 1000)

            # 记录 OCR 请求统计
            ocr_service_monitor.record_request(
                pages=1,
                chars=len(full_text),
                processing_time_ms=processing_time_ms,
                success=True,
            )

            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "image",
                "parse_method": "ocr",
                "is_scanned": True,
                "total_pages": 1,
                "text_content": full_text,
                "pages": [{
                    "page": 1,
                    "text": full_text,
                    "ocr_lines": ocr_result.get("lines", []),
                }],
                "charts": [],
                "chart_count": 0,
                "char_count": len(full_text),
                "ocr_result": ocr_result,
            }

            logger.info(f"图片 OCR 解析成功: {file_path}, 耗时 {processing_time_ms}ms")
            return result

        except Exception as e:
            processing_time_ms = int((time.time() - start_time) * 1000)
            ocr_service_monitor.record_request(
                pages=1,
                chars=0,
                processing_time_ms=processing_time_ms,
                success=False,
            )
            logger.error(f"图片 OCR 解析失败: {e}")
            return None

    def _parse_pdf(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        try:
            import fitz
            doc = fitz.open(file_path)
            if doc.is_encrypted:
                doc.close()
                return None, "PDF文件已加密"

            pages_text = []
            charts_data = []
            total_pages = len(doc)

            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    pages_text.append({
                        "page": page_num + 1,
                        "text": text.strip(),
                    })

                images = page.get_images(full=True)
                for img_idx, img_info in enumerate(images):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image_data = base_image["image"]
                            if len(image_data) > 100:
                                charts_data.append({
                                    "page": page_num + 1,
                                    "image_index": img_idx,
                                    "size": len(image_data),
                                    "ext": base_image.get("ext", "png"),
                                    "description": f"第{page_num + 1}页图片{img_idx + 1}",
                                })
                    except Exception:
                        pass

            doc.close()

            full_text = "\n".join([p["text"] for p in pages_text])
            full_text = self._clean_text(full_text)

            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "pdf",
                "parse_method": "direct",
                "is_scanned": False,
                "total_pages": total_pages,
                "text_content": full_text,
                "pages": pages_text,
                "charts": charts_data,
                "chart_count": len(charts_data),
                "char_count": len(full_text),
            }
            logger.info(f"PDF解析成功: {file_path}, 共{total_pages}页, {len(charts_data)}张图片")
            return result, "PDF解析成功"
        except ImportError:
            return None, "缺少PyMuPDF依赖"
        except Exception as e:
            return None, f"PDF解析失败: {str(e)}"

    def _parse_word(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        ext = self._get_extension(file_path)

        # .docx 使用 python-docx 解析
        if ext == "docx":
            return self._parse_docx(file_path)
        # .doc 旧版二进制格式，使用多种策略解析
        else:
            return self._parse_doc(file_path)

    def _parse_docx(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        """解析 .docx 文件 (python-docx)"""
        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    tables_data.append({
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "data": table_data,
                        "description": f"表格{table_idx + 1}",
                    })

            full_text = "\n".join(paragraphs)
            full_text = self._clean_text(full_text)

            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "docx",
                "parse_method": "direct",
                "is_scanned": False,
                "total_paragraphs": len(paragraphs),
                "text_content": full_text,
                "paragraphs": paragraphs,
                "tables": tables_data,
                "table_count": len(tables_data),
                "char_count": len(full_text),
            }
            logger.info(f"Word解析成功: {file_path}, {len(paragraphs)}段落, {len(tables_data)}表格")
            return result, "Word解析成功"
        except ImportError:
            return None, "缺少python-docx依赖"
        except Exception as e:
            return None, f"Word解析失败: {str(e)}"

    def _parse_doc(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        """
        解析 .doc 旧版二进制格式

        策略 (按优先级):
        1. antiword 命令行工具 (如果已安装)
        2. catdoc 命令行工具 (如果已安装)
        3. libreoffice 转换为 docx 后解析
        4. 降级: 二进制提取可读文本
        """
        file_name = os.path.basename(file_path)
        text = None
        parse_method = ""

        # 策略1: antiword
        text, parse_method = self._try_antiword(file_path)
        if text:
            logger.info(f".doc解析(antiword): {file_path}, {len(text)}字符")
        else:
            # 策略2: catdoc
            text, parse_method = self._try_catdoc(file_path)
            if text:
                logger.info(f".doc解析(catdoc): {file_path}, {len(text)}字符")
            else:
                # 策略3: libreoffice 转换
                text, parse_method = self._try_libreoffice_convert(file_path)
                if text:
                    logger.info(f".doc解析(libreoffice): {file_path}, {len(text)}字符")
                else:
                    # 策略4: 降级二进制提取
                    text, parse_method = self._extract_text_from_binary(file_path)
                    if text:
                        logger.info(f".doc解析(二进制提取): {file_path}, {len(text)}字符")

        if not text:
            return None, ".doc文件解析失败，请安装 antiword (apt-get install antiword) 或 catdoc (apt-get install catdoc) 或 libreoffice"

        full_text = self._clean_text(text)

        result = {
            "file_name": file_name,
            "file_format": "doc",
            "parse_method": parse_method,
            "is_scanned": False,
            "total_paragraphs": len([p for p in full_text.split("\n") if p.strip()]),
            "text_content": full_text,
            "paragraphs": [p.strip() for p in full_text.split("\n") if p.strip()],
            "tables": [],
            "table_count": 0,
            "char_count": len(full_text),
        }
        return result, "Word(.doc)解析成功"

    def _try_antiword(self, file_path: str) -> Tuple[Optional[str], str]:
        """使用 antiword 命令解析 .doc"""
        import shutil as shutil_mod
        if not shutil_mod.which("antiword"):
            return None, ""
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", "-m", "UTF-8.txt", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip(), "antiword"
        except Exception as e:
            logger.debug(f"antiword 解析失败: {e}")
        return None, ""

    def _try_catdoc(self, file_path: str) -> Tuple[Optional[str], str]:
        """使用 catdoc 命令解析 .doc"""
        import shutil as shutil_mod
        if not shutil_mod.which("catdoc"):
            return None, ""
        try:
            import subprocess
            result = subprocess.run(
                ["catdoc", "-d", "utf-8", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip(), "catdoc"
        except Exception as e:
            logger.debug(f"catdoc 解析失败: {e}")
        return None, ""

    def _try_libreoffice_convert(self, file_path: str) -> Tuple[Optional[str], str]:
        """使用 LibreOffice 转换 .doc 为 .docx 后解析"""
        import shutil as shutil_mod
        if not shutil_mod.which("libreoffice") and not shutil_mod.which("soffice"):
            return None, ""
        try:
            import subprocess
            import tempfile
            convert_dir = tempfile.mkdtemp()
            soffice = shutil_mod.which("libreoffice") or shutil_mod.which("soffice")
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", convert_dir, file_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                converted_path = os.path.join(convert_dir, base_name + ".docx")
                if os.path.exists(converted_path):
                    parsed, msg = self._parse_docx(converted_path)
                    if parsed:
                        text = parsed.get("text_content", "")
                        # 清理临时文件
                        try:
                            shutil_mod.rmtree(convert_dir)
                        except Exception:
                            pass
                        return text, "libreoffice-convert"
        except Exception as e:
            logger.debug(f"LibreOffice 转换失败: {e}")
        return None, ""

    def _extract_text_from_binary(self, file_path: str) -> Tuple[Optional[str], str]:
        """
        降级策略: 从 .doc 二进制文件中提取可读文本
        基于 OLE 复合文档格式，提取 WordDocument 流中的文本
        """
        try:
            with open(file_path, "rb") as f:
                raw = f.read()

            # 尝试使用 olefile 解析 OLE 结构
            try:
                import olefile
                ole = olefile.OleFileIO(file_path)
                if ole.exists('WordDocument'):
                    stream = ole.openstream('WordDocument')
                    word_stream = stream.read()
                    ole.close()

                    # 从 WordDocument 流中提取 UTF-16LE 文本
                    text = self._extract_utf16_text(word_stream)
                    if text and len(text) > 20:
                        return text, "binary-ole"
            except ImportError:
                pass
            except Exception as e:
                logger.debug(f"OLE 解析失败: {e}")

            # 最终降级: 提取所有可打印字符
            text = self._extract_utf16_text(raw)
            if text and len(text) > 20:
                return text, "binary-raw"

            return None, ""
        except Exception as e:
            logger.debug(f"二进制文本提取失败: {e}")
            return None, ""

    @staticmethod
    def _extract_utf16_text(data: bytes) -> Optional[str]:
        """从二进制数据中提取 UTF-16LE 编码的文本"""
        try:
            # 尝试 UTF-16LE 解码
            text = data.decode('utf-16-le', errors='ignore')
            # 过滤出可打印字符
            import re as re_mod
            # 保留中文、英文、数字、常见标点
            cleaned = re_mod.sub(r'[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffefa-zA-Z0-9\s\.\,\;\:\!\?\-\(\)\[\]\{\}\'\"\/\\@#\$%\^&\*\+\=\-_~`<>|]', '', text)
            # 压缩多余空白
            cleaned = re_mod.sub(r'\s{3,}', '\n', cleaned)
            cleaned = cleaned.strip()
            return cleaned if len(cleaned) > 20 else None
        except Exception:
            return None

    def _parse_ofd(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        """
        解析 OFD 版式文档

        使用项目内置的 ofd_service.OFDParser 进行解析：
        - 直接提取文本内容（parse_file）
        - 若文本提取为空，尝试转 PDF 后使用 PDF 解析
        """
        try:
            from ofd_service import OFDParser

            parser = OFDParser()
            ofd_result = parser.parse_file(file_path)

            if not ofd_result.valid:
                error_msg = ofd_result.error or "OFD 文件格式无效"
                logger.warning(f"OFD 解析失败: {error_msg}")
                return None, f"OFD 解析失败: {error_msg}"

            full_text = ofd_result.text or ""
            full_text = self._clean_text(full_text)

            pages_text = []
            for page_info in ofd_result.pages:
                page_text = page_info.get("text", "").strip()
                if page_text:
                    pages_text.append({
                        "page": page_info.get("page", 0),
                        "text": page_text,
                    })

            charts_data = []
            if ofd_result.has_images and ofd_result.image_resources:
                for img_idx, img_info in enumerate(ofd_result.image_resources):
                    charts_data.append({
                        "page": 0,
                        "image_index": img_idx,
                        "size": img_info.get("size", 0),
                        "ext": img_info.get("format", "PNG").lower(),
                        "description": f"OFD内嵌图片{img_idx + 1}",
                    })

            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "ofd",
                "parse_method": "ofd_parser",
                "is_scanned": False,
                "total_pages": ofd_result.page_count,
                "text_content": full_text,
                "pages": pages_text,
                "charts": charts_data,
                "chart_count": len(charts_data),
                "char_count": len(full_text),
                "metadata": ofd_result.metadata,
                "ofd_info": {
                    "valid": ofd_result.valid,
                    "page_count": ofd_result.page_count,
                    "has_images": ofd_result.has_images,
                    "has_fonts": ofd_result.has_fonts,
                    "image_count": len(ofd_result.image_resources),
                    "parse_elapsed": ofd_result.elapsed,
                },
            }

            logger.info(
                f"OFD解析成功: {file_path}, 共{ofd_result.page_count}页, "
                f"{len(full_text)}字符, {len(charts_data)}张图片, 耗时{ofd_result.elapsed}s"
            )
            return result, "OFD解析成功"

        except ImportError:
            logger.error("缺少 ofd_service 或 easyofd 依赖")
            return None, "缺少OFD解析依赖(easyofd)，请运行 pip install easyofd"
        except Exception as e:
            logger.error(f"OFD解析失败: {e}")
            return None, f"OFD解析失败: {str(e)}"

    def _parse_txt(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            text = self._clean_text(text)
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "txt",
                "parse_method": "direct",
                "is_scanned": False,
                "text_content": text,
                "paragraphs": paragraphs,
                "tables": [],
                "table_count": 0,
                "chart_count": 0,
                "char_count": len(text),
            }
            return result, "TXT解析成功"
        except Exception as e:
            return None, f"TXT解析失败: {str(e)}"

    def _parse_markdown(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], str]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            text = self._clean_text(text)
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            result = {
                "file_name": os.path.basename(file_path),
                "file_format": "md",
                "parse_method": "direct",
                "is_scanned": False,
                "text_content": text,
                "paragraphs": paragraphs,
                "tables": [],
                "table_count": 0,
                "chart_count": 0,
                "char_count": len(text),
            }
            return result, "Markdown解析成功"
        except Exception as e:
            return None, f"Markdown解析失败: {str(e)}"

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r' *\n *', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()
        return text

    def _make_chunk_id(self, file_id: int, chunk_index: int) -> str:
        raw = f"{file_id}_{chunk_index}"
        return str(uuid.uuid5(uuid.NAMESPACE_DNS, raw))

    def chunk_text(self, parsed_data: Dict[str, Any], file_id: int,
                   chunk_size: int = None, chunk_overlap: int = None,
                   chunk_separator: str = None) -> List[Dict[str, Any]]:
        chunks = []
        text = parsed_data.get("text_content", "")
        file_name = parsed_data.get("file_name", "")

        if not text:
            return chunks

        _chunk_size = chunk_size if chunk_size is not None else self.chunk_size
        _chunk_overlap = chunk_overlap if chunk_overlap is not None else self.chunk_overlap
        _separator = chunk_separator if chunk_separator else None

        if _separator:
            # 按自定义分隔符预分割, 再按 chunk_size 切片
            raw_parts = re.split(re.escape(_separator), text)
            parts = [p.strip() for p in raw_parts if p.strip()]
            for part in parts:
                part_chunks = self._chunk_by_sliding_window(
                    part, file_id, _chunk_size, _chunk_overlap, len(chunks)
                )
                chunks.extend(part_chunks)
        else:
            # 默认: 直接按 chunk_size 固定长度切片 (滑动窗口)
            chunks = self._chunk_by_sliding_window(
                text, file_id, _chunk_size, _chunk_overlap, 0
            )

        # 附加表格
        tables = parsed_data.get("tables", [])
        for table_idx, table in enumerate(tables):
            table_text = f"表格{table_idx + 1}: {json.dumps(table.get('data', []), ensure_ascii=False)}"
            chunk_id = self._make_chunk_id(file_id, len(chunks))
            chunks.append({
                "chunk_id": chunk_id,
                "text": table_text,
                "index": len(chunks),
                "file_name": file_name,
                "file_id": file_id,
                "is_table": True,
                "table_index": table_idx,
            })

        # 附加图片/图表
        charts = parsed_data.get("charts", [])
        for chart_idx, chart in enumerate(charts):
            chart_text = f"图片/图表{chart_idx + 1}: {chart.get('description', '')}"
            chunk_id = self._make_chunk_id(file_id, len(chunks))
            chunks.append({
                "chunk_id": chunk_id,
                "text": chart_text,
                "index": len(chunks),
                "file_name": file_name,
                "file_id": file_id,
                "is_chart": True,
                "chart_index": chart_idx,
            })

        sep_info = f", separator='{_separator}'" if _separator else ""
        logger.info(f"文件 {file_name} 切片完成: {len(chunks)} 个切片 (size={_chunk_size}, overlap={_chunk_overlap}{sep_info})")
        return chunks

    def _chunk_by_sliding_window(self, text: str, file_id: int,
                                  chunk_size: int, chunk_overlap: int,
                                  start_index: int = 0) -> List[Dict[str, Any]]:
        """
        使用固定长度滑动窗口对文本进行切片

        Args:
            text: 待切片的文本
            file_id: 文件ID
            chunk_size: 每个切片的最大字符数
            chunk_overlap: 相邻切片的重叠字符数
            start_index: 起始切片索引

        Returns:
            切片列表
        """
        chunks = []
        chunk_index = start_index
        file_name = ""

        if not text:
            return chunks

        # 计算有效步长 (chunk_size - chunk_overlap)
        step = max(1, chunk_size - chunk_overlap)
        text_len = len(text)

        start = 0
        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunk_text_content = text[start:end].strip()

            if chunk_text_content:
                chunk_id = self._make_chunk_id(file_id, chunk_index)
                chunks.append({
                    "chunk_id": chunk_id,
                    "text": chunk_text_content,
                    "index": chunk_index,
                    "file_name": file_name,
                    "file_id": file_id,
                })
                chunk_index += 1

            # 移动窗口
            if end >= text_len:
                break
            start = start + step

        return chunks


document_parser = DocumentParser()
